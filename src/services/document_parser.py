"""
Clean Document Parser Service.
Simplified and readable document parsing with proper error handling.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any, List
import hashlib

from pypdf import PdfReader
from docx import Document as DocxDocument
import markdown

from src.core.exceptions import DocumentParsingError
from src.utils.metadata import protect_class
from config.settings import settings


class BaseParser(ABC):
    """Base parser interface."""

    @abstractmethod
    def parse(self, file_path: Path) -> str:
        """Parse document and extract text."""

    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""


@protect_class
class PDFParser(BaseParser):
    """PDF document parser."""

    def parse(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(file_path)
            text_content = []

            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"[Page {page_num}]\n{page_text}")

            if not text_content:
                raise DocumentParsingError(
                    "No text content found in PDF", file_path=str(file_path)
                )

            return "\n\n".join(text_content)

        except Exception as e:
            if isinstance(e, DocumentParsingError):
                raise
            raise DocumentParsingError(
                f"Failed to parse PDF: {str(e)}", file_path=str(file_path)
            ) from e

    def get_supported_extensions(self) -> List[str]:
        return [".pdf"]


@protect_class
class WordParser(BaseParser):
    """Word document parser."""

    def parse(self, file_path: Path) -> str:
        """Extract text from Word document."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract table content
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text.strip():
                    text_parts.append(table_text)

            if not text_parts:
                raise DocumentParsingError(
                    "No text content found in Word document", file_path=str(file_path)
                )

            return "\n\n".join(text_parts)

        except Exception as e:
            if isinstance(e, DocumentParsingError):
                raise
            raise DocumentParsingError(
                f"Failed to parse Word document: {str(e)}", file_path=str(file_path)
            ) from e

    def _extract_table_text(self, table) -> str:
        """Extract text from a Word table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows)

    def get_supported_extensions(self) -> List[str]:
        return [".docx"]


@protect_class
class MarkdownParser(BaseParser):
    """Markdown document parser."""

    def parse(self, file_path: Path) -> str:
        """Extract and convert markdown to plain text."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                md_content = f.read()

            if not md_content.strip():
                raise DocumentParsingError(
                    "Empty markdown file", file_path=str(file_path)
                )

            # Convert markdown to HTML then extract text
            html = markdown.markdown(md_content)

            # Simple HTML tag removal
            import re

            text = re.sub("<[^<]+?>", "", html)
            text = text.replace("&nbsp;", " ").replace("&lt;", "<").replace("&gt;", ">")

            return text.strip()

        except Exception as e:
            if isinstance(e, DocumentParsingError):
                raise
            raise DocumentParsingError(
                f"Failed to parse Markdown: {str(e)}", file_path=str(file_path)
            ) from e

    def get_supported_extensions(self) -> List[str]:
        return [".md", ".markdown"]


@protect_class
class TextParser(BaseParser):
    """Plain text parser."""

    def parse(self, file_path: Path) -> str:
        """Read plain text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            if not content.strip():
                raise DocumentParsingError("Empty text file", file_path=str(file_path))

            return content.strip()

        except Exception as e:
            if isinstance(e, DocumentParsingError):
                raise
            raise DocumentParsingError(
                f"Failed to parse text file: {str(e)}", file_path=str(file_path)
            ) from e

    def get_supported_extensions(self) -> List[str]:
        return [".txt"]


@protect_class
class DocumentParserService:
    """Main document parser service."""

    def __init__(self):
        self.parsers = self._initialize_parsers()
        self._extension_map = self._build_extension_map()

    def _initialize_parsers(self) -> List[BaseParser]:
        """Initialize all available parsers."""
        return [
            PDFParser(),
            WordParser(),
            MarkdownParser(),
            TextParser(),
        ]

    def _build_extension_map(self) -> Dict[str, BaseParser]:
        """Build mapping from file extensions to parsers."""
        extension_map = {}
        for parser in self.parsers:
            for ext in parser.get_supported_extensions():
                extension_map[ext.lower()] = parser
        return extension_map

    def is_supported(self, file_path: str) -> bool:
        """Check if file type is supported."""
        extension = Path(file_path).suffix.lower()
        return extension in self._extension_map

    def get_supported_extensions(self) -> List[str]:
        """Get all supported file extensions."""
        return list(self._extension_map.keys())

    def validate_file(self, file_path: str) -> None:
        """Validate file before parsing."""
        path = Path(file_path)

        if not path.exists():
            raise DocumentParsingError(
                f"File not found: {file_path}", file_path=file_path
            )

        if not self.is_supported(file_path):
            supported = ", ".join(self.get_supported_extensions())
            raise DocumentParsingError(
                f"Unsupported file type. Supported: {supported}", file_path=file_path
            )

        file_size = path.stat().st_size
        if file_size > settings.api.max_file_size:
            raise DocumentParsingError(
                f"File too large: {file_size} bytes. Max: {settings.api.max_file_size} bytes",
                file_path=file_path,
            )

    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """Parse document and return extracted content with metadata."""
        # Validate file first
        self.validate_file(file_path)

        path = Path(file_path)
        extension = path.suffix.lower()
        parser = self._extension_map[extension]

        # Parse document
        content = parser.parse(path)

        # Calculate content hash
        content_hash = hashlib.sha256(content.encode()).hexdigest()

        return {
            "file_name": path.name,
            "file_path": str(path),
            "file_type": extension,
            "content": content,
            "content_hash": content_hash,
            "char_count": len(content),
            "word_count": len(content.split()),
            "metadata": {
                "parser": parser.__class__.__name__,
                "file_size": path.stat().st_size,
            },
        }


# Global instance
document_parser = DocumentParserService()
